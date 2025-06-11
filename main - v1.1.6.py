import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import re
import threading
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import math
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class CodeFormatterApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        ctk.set_appearance_mode("light")

        self.title("软件著作权代码格式化工具-陈安锦-v1.1.6")
        self.geometry("800x750")

        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(8, weight=1)

        self.create_widgets()

    def create_widgets(self):
        # --- 源码选择 ---
        source_frame = ctk.CTkFrame(self)
        source_frame.grid(row=0, column=0, padx=10, pady=5, sticky="ew")
        source_frame.grid_columnconfigure(1, weight=1)

        ctk.CTkLabel(source_frame, text="源码路径:").grid(row=0, column=0, padx=5, pady=5)
        self.source_path_entry = ctk.CTkEntry(source_frame)
        self.source_path_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        ctk.CTkButton(source_frame, text="选择文件", command=self.select_source_file).grid(row=0, column=2, padx=5, pady=5)
        ctk.CTkButton(source_frame, text="选择文件夹", command=self.select_source_folder).grid(row=0, column=3, padx=5, pady=5)

        # --- 代码类型 ---
        lang_frame = ctk.CTkFrame(self)
        lang_frame.grid(row=1, column=0, padx=10, pady=5, sticky="ew")
        ctk.CTkLabel(lang_frame, text="代码类型:").pack(side="left", padx=5, pady=5)
        
        self.lang_var = tk.StringVar(value="Auto")
        languages = [("自动检测", "Auto"), ("Java", "Java"), ("Python", "Python"), ("C/C++", "C"), ("C#", "CS"), ("JavaScript", "JS")]
        for lang, val in languages:
            ctk.CTkRadioButton(lang_frame, text=lang, variable=self.lang_var, value=val).pack(side="left", padx=5, pady=5)

        # --- 文档信息设置 ---
        doc_info_frame = ctk.CTkFrame(self)
        doc_info_frame.grid(row=2, column=0, padx=10, pady=5, sticky="ew")

        ctk.CTkLabel(doc_info_frame, text="软件名称:").pack(side="left", padx=5, pady=5)
        self.software_name_entry = ctk.CTkEntry(doc_info_frame)
        self.software_name_entry.insert(0, "example")
        self.software_name_entry.pack(side="left", padx=5, pady=5)

        ctk.CTkLabel(doc_info_frame, text="版本号:").pack(side="left", padx=5, pady=5)
        self.version_entry = ctk.CTkEntry(doc_info_frame)
        self.version_entry.insert(0, "V1.0")
        self.version_entry.pack(side="left", padx=5, pady=5)

        # --- 代码处理选项 ---
        process_options_frame = ctk.CTkFrame(self)
        process_options_frame.grid(row=3, column=0, padx=10, pady=5, sticky="ew")
        ctk.CTkLabel(process_options_frame, text="代码处理选项:").pack(side="left", padx=5)

        self.remove_comments_var = ctk.BooleanVar(value=True)
        ctk.CTkCheckBox(process_options_frame, text="移除注释", variable=self.remove_comments_var).pack(side="left", padx=5)
        self.remove_blank_lines_var = ctk.BooleanVar(value=True)
        ctk.CTkCheckBox(process_options_frame, text="清除空行", variable=self.remove_blank_lines_var).pack(side="left", padx=5)

        # --- 格式设置 ---
        format_frame = ctk.CTkFrame(self)
        format_frame.grid(row=4, column=0, padx=10, pady=5, sticky="ew")
        
        ctk.CTkLabel(format_frame, text="字体:").pack(side="left", padx=5, pady=5)
        self.font_name_combo = ctk.CTkComboBox(format_frame, values=["Times New Roman", "宋体", "Courier New"])
        self.font_name_combo.set("Times New Roman")
        self.font_name_combo.pack(side="left", padx=5, pady=5)

        ctk.CTkLabel(format_frame, text="字号:").pack(side="left", padx=5, pady=5)
        self.font_size_combo = ctk.CTkComboBox(format_frame, values=[str(i) for i in range(8, 17)])
        self.font_size_combo.set("10.5")
        self.font_size_combo.pack(side="left", padx=5, pady=5)
        
        ctk.CTkLabel(format_frame, text="每页行数:").pack(side="left", padx=5, pady=5)
        self.lines_per_page_combo = ctk.CTkComboBox(format_frame, values=[str(i) for i in [30, 40, 50, 60, 70]])
        self.lines_per_page_combo.set("50")
        self.lines_per_page_combo.pack(side="left", padx=5, pady=5)

        # --- 页面范围 ---
        page_range_frame = ctk.CTkFrame(self)
        page_range_frame.grid(row=5, column=0, padx=10, pady=5, sticky="ew")
        self.page_range_var = tk.StringVar(value="first_last_30")
        ctk.CTkRadioButton(page_range_frame, text="提取前后各30页", variable=self.page_range_var, value="first_last_30").pack(side="left", padx=5, pady=5)
        ctk.CTkRadioButton(page_range_frame, text="提取全部页面", variable=self.page_range_var, value="all").pack(side="left", padx=5, pady=5)

        # --- 输出 ---
        output_frame = ctk.CTkFrame(self)
        output_frame.grid(row=6, column=0, padx=10, pady=5, sticky="ew")
        output_frame.grid_columnconfigure(1, weight=1)
        
        ctk.CTkLabel(output_frame, text="输出目录:").grid(row=0, column=0, padx=5, pady=5)
        self.output_dir_entry = ctk.CTkEntry(output_frame)
        self.output_dir_entry.insert(0, "formatted_code")
        self.output_dir_entry.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        ctk.CTkButton(output_frame, text="选择目录", command=self.select_output_dir).grid(row=0, column=2, padx=5, pady=5)

        # --- 控制按钮 ---
        action_frame = ctk.CTkFrame(self)
        action_frame.grid(row=7, column=0, padx=10, pady=10, sticky="ew")
        self.start_button = ctk.CTkButton(action_frame, text="开始处理", command=self.start_processing_wrapper)
        self.start_button.pack(pady=5)

        # --- 进度和日志 ---
        progress_log_frame = ctk.CTkFrame(self)
        progress_log_frame.grid(row=8, column=0, padx=10, pady=5, sticky="nsew")
        progress_log_frame.grid_rowconfigure(1, weight=1)
        progress_log_frame.grid_columnconfigure(0, weight=1)

        self.progress_bar = ctk.CTkProgressBar(progress_log_frame)
        self.progress_bar.set(0)
        self.progress_bar.grid(row=0, column=0, padx=5, pady=5, sticky="ew")

        self.log_textbox = ctk.CTkTextbox(progress_log_frame, state="disabled")
        self.log_textbox.grid(row=1, column=0, padx=5, pady=5, sticky="nsew")

    def select_source_file(self):
        path = filedialog.askopenfilename()
        if path:
            self.source_path_entry.delete(0, tk.END)
            self.source_path_entry.insert(0, path)

    def select_source_folder(self):
        path = filedialog.askdirectory()
        if path:
            self.source_path_entry.delete(0, tk.END)
            self.source_path_entry.insert(0, path)

    def select_output_dir(self):
        path = filedialog.askdirectory()
        if path:
            self.output_dir_entry.delete(0, tk.END)
            self.output_dir_entry.insert(0, path)

    def log(self, message):
        def _log():
            self.log_textbox.configure(state="normal")
            self.log_textbox.insert(tk.END, message + "\n")
            self.log_textbox.configure(state="disabled")
            self.log_textbox.see(tk.END)
        self.after(0, _log)

    def update_progress(self, value):
        def _update():
            self.progress_bar.set(value)
        self.after(0, _update)

    def start_processing_wrapper(self):
        self.start_button.configure(state="disabled", text="处理中...")
        self.progress_bar.set(0)
        self.log_textbox.configure(state="normal")
        self.log_textbox.delete("1.0", tk.END)
        self.log_textbox.configure(state="disabled")

        thread = threading.Thread(target=self.process_files)
        thread.start()

    def process_files(self):
        try:
            config = self.get_config()
            if not config["source_path"] or not config["output_dir"]:
                messagebox.showerror("错误", "源码路径和输出目录不能为空")
                return

            self.log("开始处理...")
            self.update_progress(0.1)

            source_files = self.find_source_files(config["source_path"], config["lang"])
            if not source_files:
                self.log("错误：未找到指定的代码文件。")
                return

            self.log(f"找到 {len(source_files)} 个代码文件。")
            all_lines = self.read_and_process_files(source_files, config)
            
            self.log("代码读取和处理完成。")
            self.update_progress(0.6)

            self.create_word_document(all_lines, config)
            self.log("Word文档生成成功！")
            
            output_filename = f'{config["software_name"]}_{config["version"]}_源代码.docx'
            self.log(f"文件保存在: {os.path.join(config['output_dir'], output_filename)}")

        except Exception as e:
            self.log(f"发生错误: {e}")
            messagebox.showerror("错误", f"处理过程中发生错误:\n{e}")
        finally:
            def _finalize():
                self.start_button.configure(state="normal", text="开始处理")
                self.update_progress(1.0)
            self.after(0, _finalize)

    def get_config(self):
        return {
            "source_path": self.source_path_entry.get(),
            "lang": self.lang_var.get(),
            "software_name": self.software_name_entry.get(),
            "version": self.version_entry.get(),
            "remove_comments": self.remove_comments_var.get(),
            "remove_blank_lines": self.remove_blank_lines_var.get(),
            "font_name": self.font_name_combo.get(),
            "font_size": float(self.font_size_combo.get()),
            "lines_per_page": int(self.lines_per_page_combo.get()),
            "page_range": self.page_range_var.get(),
            "output_dir": self.output_dir_entry.get(),
        }

    def find_source_files(self, path, lang):
        extensions = {
            "Java": [".java"],
            "Python": [".py"],
            "C": [".c", ".cpp", ".h", ".hpp"],
            "CS": [".cs"],
            "JS": [".js"],
            "Auto": [".java", ".py", ".c", ".cpp", ".h", ".hpp", ".cs", ".js"],
        }
        
        target_exts = extensions.get(lang, [])
        source_files = []
        if os.path.isfile(path):
            if any(path.endswith(ext) for ext in target_exts):
                source_files.append(path)
        elif os.path.isdir(path):
            for root, _, files in os.walk(path):
                for file in files:
                    if any(file.endswith(ext) for ext in target_exts):
                        source_files.append(os.path.join(root, file))
        return source_files

    def read_and_process_files(self, files, config):
        all_lines = []
        total_files = len(files)
        for i, file_path in enumerate(files):
            self.log(f"处理文件: {file_path}")
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                
                if config["remove_comments"]:
                    # Simple comment removal, can be improved
                    if config["lang"] in ["Java", "C", "CS", "JS"]:
                        content = re.sub(r'//.*', '', content)
                        content = re.sub(r'/\*[\s\S]*?\*/', '', content)
                    elif config["lang"] == "Python":
                        content = re.sub(r'#.*', '', content)

                lines = content.split('\n')
                if config["remove_blank_lines"]:
                    lines = [line for line in lines if line.strip() != '']
                
                all_lines.extend(lines)
            except Exception as e:
                self.log(f"警告：读取或处理文件失败 {file_path}: {e}")
            self.update_progress(0.1 + 0.5 * (i + 1) / total_files)
        return all_lines

    def create_word_document(self, lines, config):
        doc = Document()

        # Set page margins
        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # --- Header Setup ---
        header = section.header
        # Clear existing default paragraph in header
        if header.paragraphs:
            header.paragraphs[0].text = ""
            # Remove any other paragraphs
            for p in header.paragraphs[1:]:
                p_element = p._element
                p_element.getparent().remove(p_element)
        
        # Add a table for layout
        table = header.add_table(rows=1, cols=3, width=section.page_width - section.left_margin - section.right_margin)
        
        # Center cell for software name and version
        p_center = table.cell(0, 1).paragraphs[0]
        p_center.text = f'{config["software_name"]} {config["version"]}'
        p_center.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Right cell for page number
        p_right = table.cell(0, 2).paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add page number field using OOXML
        run = p_right.add_run()
        fldChar_begin = OxmlElement('w:fldChar')
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar_end = OxmlElement('w:fldChar')
        fldChar_end.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar_begin)
        run._r.append(instrText)
        run._r.append(fldChar_end)

        lines_per_page = config["lines_per_page"]
        total_lines = len(lines)
        total_pages = math.ceil(total_lines / lines_per_page)
        self.log(f"--- 开始生成文档 ---")
        self.log(f"代码总行数: {total_lines}, 每页行数: {lines_per_page}")
        self.log(f"计算出的总页数: {total_pages}")

        pages_to_extract = []
        if config["page_range"] == "all":
            self.log("选项：提取全部页面。")
            pages_to_extract = list(range(total_pages))
        else:  # 'first_last_30'
            self.log("选项：提取前后各30页。")
            if total_pages <= 60:
                self.log(f"总页数 ({total_pages}) 不超过60页，将提取全部页面。")
                pages_to_extract = list(range(total_pages))
            else:
                self.log(f"总页数 ({total_pages}) 超过60页，将提取前后各30页。")
                pages_to_extract.extend(range(30))  # First 30 pages
                pages_to_extract.extend(range(total_pages - 30, total_pages))  # Last 30 pages
        
        pages_to_extract = sorted(list(set(pages_to_extract))) # Just in case of any overlap logic change

        self.log(f"最终选择提取的页数: {len(pages_to_extract)}")
        # Log the first and last few pages to verify, to avoid flooding the log
        if len(pages_to_extract) > 20:
            self.log(f"提取页面列表 (示例): {pages_to_extract[:10]}...{pages_to_extract[-10:]}")
        else:
            self.log(f"提取页面列表: {pages_to_extract}")

        for i, page_num in enumerate(pages_to_extract):
            start_line = page_num * lines_per_page
            end_line = start_line + lines_per_page
            page_lines = lines[start_line:end_line]
            
            # Reworked logic: add each line as a separate, formatted paragraph
            # to have finer control over spacing and prevent content overflow.
            for line_content in page_lines:
                p = doc.add_paragraph()
                
                pf = p.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0
                
                run = p.add_run(line_content)
                font = run.font
                font.name = config["font_name"]
                font.size = Pt(config["font_size"])
                font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            
            if i < len(pages_to_extract) - 1:
                doc.add_page_break()
        
        output_dir = config['output_dir']
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        output_filename = f'{config["software_name"]}_{config["version"]}_源代码.docx'
        doc.save(os.path.join(output_dir, output_filename))


if __name__ == "__main__":
    app = CodeFormatterApp()
    app.mainloop() 